"""
Парсинг документов (TXT, PDF, DOCX)
"""
import os
import logging
from pathlib import Path
from typing import Optional

import pypdf
from docx import Document

logger = logging.getLogger(__name__)


async def extract_text_from_file(file_path: Path, file_name: str) -> str:
    """Извлекает текст из TXT, PDF или DOCX."""
    extracted_text = ""
    try:
        if file_name.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read()
        elif file_name.lower().endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text:
                        extracted_text += text + "\n"
        elif file_name.lower().endswith('.docx'):
            word_doc = Document(file_path)
            for para in word_doc.paragraphs:
                extracted_text += para.text + "\n"
    except Exception as e:
        logger.error(f"Ошибка парсинга файла {file_name}: {e}")
    return extracted_text


async def extract_document_text(context, doc) -> Optional[str]:
    """Скачивает и извлекает текст из документа для reply-анализа."""
    safe_name = Path(doc.file_name).name if doc.file_name else "unknown"
    file_path = Path("temp") / f"temp_extract_{os.urandom(4).hex()}_{safe_name}"
    file_path.parent.mkdir(exist_ok=True)
    try:
        file = await context.bot.get_file(doc.file_id)
        await file.download_to_drive(file_path)
        
        extracted_text = ""
        if doc.file_name.lower().endswith('.txt'):
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                extracted_text = f.read()
        elif doc.file_name.lower().endswith('.pdf'):
            with open(file_path, 'rb') as f:
                reader = pypdf.PdfReader(f)
                for page in reader.pages:
                    text = page.extract_text()
                    if text: extracted_text += text + "\n"
        elif doc.file_name.lower().endswith('.docx'):
            word_doc = Document(file_path)
            for para in word_doc.paragraphs:
                extracted_text += para.text + "\n"
        
        if len(extracted_text) > 30000:
            extracted_text = extracted_text[:30000] + "\n...[текст обрезан]..."
        return extracted_text.strip() if extracted_text.strip() else None
    except Exception as e:
        logger.error(f"Ошибка при извлечении текста из {doc.file_name}: {e}")
        return None
    finally:
        if file_path.exists():
            file_path.unlink()
